import sys
import requests
import traceback
from exception import customexception
from logger import logging
from llama_index.core import SimpleDirectoryReader, Document
from pdfminer.high_level import extract_text
import docx
import os

def extract_text_from_docx(file):
    """Extract text from a DOCX file"""
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def load_data(source, is_url=False, is_uploaded_file=False):
    try:
        logging.info("Data loading...")

        if is_url:
            response = requests.get(source)
            if response.status_code == 200:
                document_text = response.text
                logging.info("Data successfully fetched from URL.")
            else:
                logging.info("Failed to fetch data from URL.")
                raise customexception(f"Error fetching data from URL: {source}", traceback.format_exc())

        elif is_uploaded_file:
            file_extension = os.path.splitext(source.name)[-1].lower()
            
            if file_extension == ".txt":
                document_text = source.getvalue().decode("utf-8")
            elif file_extension == ".pdf":
                document_text = extract_text(source)
            elif file_extension == ".docx":
                document_text = extract_text_from_docx(source)
            else:
                raise customexception(f"Unsupported file format: {file_extension}", traceback.format_exc())

            logging.info("Data successfully loaded from uploaded file.")

        else:
            loader = SimpleDirectoryReader(source)
            documents = loader.load_data()
            logging.info("Data loading completed from directory.")
            return documents
        
        return [Document(text=document_text)]

    except Exception as e:
        logging.error(f"Error in loading data: {str(e)}")
        raise customexception(str(e), traceback.format_exc())


